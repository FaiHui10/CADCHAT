# -*- coding: utf-8 -*-
"""
软著申请材料生成脚本
生成软件著作权登记申请所需材料
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_software_copyright_doc():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ========================================
    # 第一部分：软件著作权登记申请表
    # ========================================
    heading = doc.add_heading('', level=1)
    run = heading.add_run('软件著作权登记申请表')
    run.font.size = Pt(18)
    run.font.bold = True
    
    # 表格1：基本信息
    table1 = doc.add_table(rows=12, cols=2)
    table1.style = 'Table Grid'
    
    data1 = [
        ('软件名称（全称）', 'CAD智能助手'),
        ('软件简称', 'CADChat'),
        ('版本号', 'V0.11'),
        ('软件分类', '应用软件'),
        ('开发完成日期', '2026-02-20'),
        ('首次发表日期', '未发表'),
        ('著作权人姓名/名称', ''),
        ('身份证号/组织机构代码', ''),
        ('通讯地址', ''),
        ('联系人', ''),
        ('联系电话', ''),
        ('电子邮箱', ''),
    ]
    
    for i, (label, value) in enumerate(data1):
        cell = table1.rows[i].cells[0]
        cell.text = label
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.width = Cm(4)
        
        cell = table1.rows[i].cells[1]
        cell.text = value
    
    doc.add_paragraph()
    
    # 申请方式
    doc.add_paragraph('申请方式：□原始取得  □继受取得')
    doc.add_paragraph('是否委托办理：□是  □否')
    
    doc.add_page_break()
    
    # ========================================
    # 第二部分：软件说明书
    # ========================================
    heading2 = doc.add_heading('', level=1)
    run2 = heading2.add_run('软件说明书')
    run2.font.size = Pt(18)
    run2.font.bold = True
    
    # 1. 软件名称
    doc.add_heading('一、软件名称', level=2)
    p = doc.add_paragraph()
    p.add_run('中文名称：').bold = True
    p.add_run('CAD智能助手')
    p = doc.add_paragraph()
    p.add_run('英文名称：').bold = True
    p.add_run('CAD Intelligent Assistant')
    p = doc.add_paragraph()
    p.add_run('软件简称：').bold = True
    p.add_run('CADChat')
    
    # 2. 软件用途
    doc.add_heading('二、软件用途', level=2)
    p = doc.add_paragraph()
    p.add_run('本软件是一款基于人工智能技术的CAD辅助工具，旨在帮助用户快速生成AutoCAD/中望CAD的LISP程序代码。')
    p = doc.add_paragraph()
    p.add_run('通过自然语言描述需求，系统可自动匹配并推荐相关命令，生成对应的LISP代码，实现CAD绘图的智能化辅助。')
    p = doc.add_paragraph()
    p.add_run('本软件采用客户端-服务端架构，支持局域网部署，可同时服务多个客户端用户。')
    p = doc.add_paragraph()
    p.add_run('核心技术包括：RAG（检索增强生成）向量检索、本地大模型推理、浏览器自动化、CAD集成等。')
    
    # 3. 硬件环境
    doc.add_heading('三、硬件环境', level=2)
    doc.add_heading('3.1 客户端硬件要求', level=3)
    doc.add_paragraph('处理器：Intel Core i5 或同等性能以上')
    doc.add_paragraph('内存：8GB 或以上')
    doc.add_paragraph('硬盘：至少 1GB 可用空间')
    doc.add_paragraph('显示器：分辨率 1920x1080 或以上')
    doc.add_paragraph('显卡：支持 OpenGL 2.0 的独立显卡（用于CAD运行）')
    
    doc.add_heading('3.2 服务端硬件要求', level=3)
    doc.add_paragraph('处理器：Intel Core i7 或同等性能以上（用于本地大模型推理）')
    doc.add_paragraph('内存：16GB 或以上（建议32GB）')
    doc.add_paragraph('硬盘：至少 10GB 可用空间（用于模型存储）')
    doc.add_paragraph('GPU：NVIDIA RTX 2060 或以上（可选，用于GPU加速推理）')
    
    # 4. 软件环境
    doc.add_heading('四、软件环境', level=2)
    doc.add_heading('4.1 客户端软件要求', level=3)
    doc.add_paragraph('操作系统：Windows 10/11 64位')
    doc.add_paragraph('运行环境：Python 3.8+')
    doc.add_paragraph('浏览器：Chrome/Edge（用于Kimi功能）')
    
    doc.add_heading('4.2 服务端软件要求', level=3)
    doc.add_paragraph('操作系统：Windows 10/11 64位')
    doc.add_paragraph('运行环境：Python 3.8+')
    doc.add_paragraph('依赖软件：Ollama（本地大模型运行平台）')
    doc.add_paragraph('CAD软件：AutoCAD 2016+ 或 中望CAD 2020+')
    
    # 5. 编程语言和版本
    doc.add_heading('五、编程语言和版本', level=2)
    doc.add_paragraph('主要编程语言：Python 3.13')
    doc.add_paragraph('前端框架：Tkinter（Python内置GUI库）')
    doc.add_paragraph('后端框架：Flask 3.0')
    doc.add_paragraph('AI模型：Ollama + Qwen3:1.7b（推理）、bge-m3（向量化）')
    doc.add_paragraph('浏览器自动化：Playwright')
    doc.add_paragraph('CAD集成：PythonWin（pywin32）')
    
    # 6. 功能描述
    doc.add_heading('六、功能描述', level=2)
    
    doc.add_heading('6.1 命令检索功能', level=3)
    doc.add_paragraph('系统提供三种命令来源的检索：')
    doc.add_paragraph('（1）基本命令：AutoCAD内置命令（如LINE、CIRCLE等）')
    doc.add_paragraph('（2）LISP命令：自定义LISP程序的功能描述')
    doc.add_paragraph('（3）用户代码：用户保存的LISP代码片段')
    doc.add_paragraph('用户输入功能需求后，系统返回TOP-3最相关的命令供用户选择。')
    
    doc.add_heading('6.2 RAG向量检索', level=3)
    doc.add_paragraph('采用RAG（检索增强生成）技术：')
    doc.add_paragraph('（1）使用bge-m3嵌入模型将命令描述向量化')
    doc.add_paragraph('（2）通过余弦相似度进行语义匹配')
    doc.add_paragraph('（3）返回TOP-3最相关结果供用户选择')
    doc.add_paragraph('（4）支持相似度阈值设置，高相似度结果跳过LLM匹配直接返回')
    
    doc.add_heading('6.3 LISP代码生成', level=3)
    doc.add_paragraph('（1）用户输入功能需求（如"画一个圆"）')
    doc.add_paragraph('（2）系统检索相关命令')
    doc.add_paragraph('（3）调用LLM生成LISP代码')
    doc.add_paragraph('（4）返回代码到客户端显示')
    doc.add_paragraph('（5）自动保存代码到本地临时文件')
    doc.add_paragraph('（6）支持代码复制到CAD中执行')
    
    doc.add_heading('6.4 用户代码管理', level=3)
    doc.add_paragraph('（1）保存用户编写的LISP代码到服务器')
    doc.add_paragraph('（2）自动提取命令名称和功能描述')
    doc.add_paragraph('（3）将用户代码纳入RAG检索范围')
    doc.add_paragraph('（4）支持用户代码的双击加载')
    doc.add_paragraph('（5）用户代码存储在独立目录，便于管理')
    
    doc.add_heading('6.5 Kimi代码分析', level=3)
    doc.add_paragraph('集成Kimi浏览器自动化：')
    doc.add_paragraph('（1）自动登录Kimi网站')
    doc.add_paragraph('（2）发送代码分析请求')
    doc.add_paragraph('（3）生成50字以内的功能描述')
    doc.add_paragraph('（4）提取分析结果用于代码保存时的自动标注')
    doc.add_paragraph('（5）支持超时处理和错误恢复')
    
    doc.add_heading('6.6 CAD集成', level=3)
    doc.add_paragraph('（1）通过PythonWin模块连接CAD')
    doc.add_paragraph('（2）实现LISP代码的自动加载执行')
    doc.add_paragraph('（3）支持AutoCAD和中望CAD双平台')
    doc.add_paragraph('（4）实时显示CAD连接状态')
    
    doc.add_heading('6.7 并发支持', level=3)
    doc.add_paragraph('（1）Flask多线程模式支持多客户端并发访问')
    doc.add_paragraph('（2）向量检索与缓存重建并行处理')
    doc.add_paragraph('（3）临时文件机制保证数据一致性')
    doc.add_paragraph('（4）请求超时保护机制')
    
    doc.add_heading('6.8 自动缓存管理', level=3)
    doc.add_paragraph('（1）使用watchdog监控命令文件变化')
    doc.add_paragraph('（2）自动触发向量缓存重建')
    doc.add_paragraph('（3）临时文件机制确保重建过程中服务不中断')
    doc.add_paragraph('（4）缓存文件格式支持numpy直接加载')
    
    doc.add_page_break()
    
    # 7. 系统架构
    doc.add_heading('七、系统架构', level=2)
    
    doc.add_heading('7.1 总体架构', level=3)
    doc.add_paragraph('本系统采用客户端-服务端（C/S）架构，整体架构分为三层：')
    doc.add_paragraph('（1）表现层：Tkinter图形用户界面')
    doc.add_paragraph('（2）业务层：Flask RESTful API服务')
    doc.add_paragraph('（3）数据层：本地向量数据库和文件系统')
    
    doc.add_heading('7.2 系统架构图', level=3)
    p = doc.add_paragraph()
    p.add_run('【图1：系统架构图】').bold = True
    try:
        doc.add_picture('dist/diagram1_system_architecture.png', width=Inches(6))
    except:
        doc.add_paragraph('（请在此处插入系统架构图）')
    
    doc.add_heading('7.3 服务端架构图', level=3)
    p = doc.add_paragraph()
    p.add_run('【图2：服务端架构图】').bold = True
    try:
        doc.add_picture('dist/diagram2_server_architecture.png', width=Inches(6))
    except:
        doc.add_paragraph('（请在此处插入服务端架构图）')
    
    doc.add_heading('7.4 客户端架构图', level=3)
    p = doc.add_paragraph()
    p.add_run('【图3：客户端架构图】').bold = True
    try:
        doc.add_picture('dist/diagram3_client_architecture.png', width=Inches(6))
    except:
        doc.add_paragraph('（请在此处插入客户端架构图）')
    
    # 8. 处理流程
    doc.add_heading('八、处理流程', level=2)
    
    doc.add_heading('8.1 整体工作流程', level=3)
    p = doc.add_paragraph()
    p.add_run('【图4：整体工作流程图】').bold = True
    try:
        doc.add_picture('dist/diagram4_workflow.png', width=Inches(6))
    except:
        doc.add_paragraph('（请在此处插入工作流程图）')
    
    doc.add_heading('8.2 命令检索流程', level=3)
    p = doc.add_paragraph()
    p.add_run('【图5：命令检索流程图】').bold = True
    try:
        doc.add_picture('dist/diagram5_retrieval_flow.png', width=Inches(5))
    except:
        doc.add_paragraph('步骤1：用户输入功能需求（如"画一个圆"）')
        doc.add_paragraph('步骤2：客户端将需求发送到服务端API')
        doc.add_paragraph('步骤3：服务端将需求文本转换为向量（bge-m3）')
        doc.add_paragraph('步骤4：计算与命令库中所有命令的余弦相似度')
        doc.add_paragraph('步骤5：排序并返回TOP-3最相似命令')
        doc.add_paragraph('步骤6：客户端显示检索结果（命令名、描述、来源）')
    
    doc.add_heading('8.3 用户代码保存流程', level=3)
    p = doc.add_paragraph()
    p.add_run('【图6：用户代码保存流程图】').bold = True
    try:
        doc.add_picture('dist/diagram6_save_flow.png', width=Inches(5))
    except:
        doc.add_paragraph('步骤1：用户在客户端点击"保存代码"')
        doc.add_paragraph('步骤2：客户端调用Kimi浏览器自动化模块')
        doc.add_paragraph('步骤3：Kimi分析代码并生成功能描述')
        doc.add_paragraph('步骤4：客户端弹出确认对话框展示描述')
        doc.add_paragraph('步骤5：用户确认后，代码保存到服务器')
        doc.add_paragraph('步骤6：服务端更新用户代码索引文件')
        doc.add_paragraph('步骤7：触发向量缓存自动重建')
    
    # 9. 界面说明
    doc.add_heading('九、界面说明', level=2)
    
    doc.add_heading('9.1 主界面布局', level=3)
    p = doc.add_paragraph()
    p.add_run('【图7：客户端主界面截图】').bold = True
    doc.add_paragraph('（请在此处插入客户端主界面截图）')
    
    doc.add_paragraph('客户端主界面分为以下区域：')
    doc.add_paragraph('（1）顶部标题栏：显示软件名称和版本')
    doc.add_paragraph('（2）功能需求输入区：用户输入想要实现的CAD功能')
    doc.add_paragraph('（3）连接状态栏：显示CAD连接状态和服务器连接状态')
    doc.add_paragraph('（4）服务器检索结果区：显示RAG检索的TOP-3结果')
    doc.add_paragraph('（5）LISP代码显示区：显示生成的LISP代码')
    doc.add_paragraph('（6）日志显示区：显示操作日志和系统信息')
    doc.add_paragraph('（7）功能按钮区：连接CAD、保存代码等操作按钮')
    
    doc.add_heading('9.2 检索结果展示', level=3)
    p = doc.add_paragraph()
    p.add_run('【图8：检索结果展示截图】').bold = True
    doc.add_paragraph('（请在此处插入检索结果展示截图）')
    
    doc.add_paragraph('检索结果以树形结构展示，每条结果包含：')
    doc.add_paragraph('（1）命令名称：如CIRCLE')
    doc.add_paragraph('（2）功能描述：如绘制圆')
    doc.add_paragraph('（3）来源类型：基本命令/LISP命令/用户代码')
    doc.add_paragraph('（4）相似度分数（如适用）')
    doc.add_paragraph('用户双击用户代码可加载完整代码内容')
    
    doc.add_heading('9.3 代码保存对话框', level=3)
    p = doc.add_paragraph()
    p.add_run('【图9：代码保存对话框截图】').bold = True
    doc.add_paragraph('（请在此处插入代码保存对话框截图）')
    
    doc.add_paragraph('保存对话框包含：')
    doc.add_paragraph('（1）命令名称输入框')
    doc.add_paragraph('（2）功能描述输入框（自动填充，可编辑）')
    doc.add_paragraph('（3）确认和取消按钮')
    
    # 10. 数据结构
    doc.add_heading('十、数据结构', level=2)
    
    doc.add_heading('10.1 命令库文件格式', level=3)
    doc.add_paragraph('基本命令库格式（autocad_basic_commands.txt）：')
    doc.add_paragraph('命令名称|功能描述|快捷键|命令类型')
    doc.add_paragraph('示例：')
    doc.add_paragraph('LINE|绘制直线|L|basic')
    doc.add_paragraph('CIRCLE|绘制圆|C|basic')
    doc.add_paragraph('RECTANG|绘制矩形|REC|basic')
    
    doc.add_heading('10.2 用户代码索引格式', level=3)
    doc.add_paragraph('用户代码索引格式（user_codes.txt）：')
    doc.add_paragraph('命令名称|功能描述|文件名|命令类型')
    doc.add_paragraph('示例：')
    doc.add_paragraph('DrawCircle|绘制圆的LISP程序|code_123456.lsp|user')
    doc.add_paragraph('DrawRect|绘制矩形的LISP程序|code_234567.lsp|user')
    
    doc.add_heading('10.3 向量缓存格式', level=3)
    doc.add_paragraph('向量缓存使用numpy数组格式存储：')
    doc.add_paragraph('（1）command_embeddings_bge_m3.npy：命令向量矩阵')
    doc.add_paragraph('（2）格式：numpy.ndarray，shape=(n_commands, embedding_dim)')
    doc.add_paragraph('（3）embedding_dim = 1024（bge-m3维度）')
    
    # 11. 创新点
    doc.add_heading('十一、创造性说明', level=2)
    
    doc.add_heading('11.1 技术创新', level=3)
    doc.add_paragraph('（1）RAG技术应用：将检索增强生成技术应用于CAD命令检索领域')
    doc.add_paragraph('    传统CAD辅助工具主要依赖关键词匹配，本系统采用向量语义检索，')
    doc.add_paragraph('    大幅提升了检索的准确性和语义理解能力。')
    
    doc.add_paragraph('（2）多源命令统一检索：实现了基本命令、LISP命令、用户代码的')
    doc.add_paragraph('    统一检索框架，用户无需切换即可搜索所有可用命令。')
    
    doc.add_paragraph('（3）浏览器自动化集成：创新性地将Kimi浏览器自动化集成到')
    doc.add_paragraph('    代码保存流程，自动生成功能描述，简化用户操作。')
    
    doc.add_paragraph('（4）自动缓存管理：使用watchdog实现文件监控，结合临时文件')
    doc.add_paragraph('    机制实现热更新，保证服务不中断。')
    
    doc.add_heading('11.2 应用创新', level=3)
    doc.add_paragraph('（1）客户端-服务端架构：支持局域网部署，可同时服务多个客户端')
    doc.add_paragraph('（2）双平台支持：同时支持AutoCAD和中望CAD')
    doc.add_paragraph('（3）本地大模型：所有AI能力本地运行，无需联网，保护数据隐私')
    doc.add_paragraph('（4）用户代码管理：用户可保存自己的LISP代码，系统自动纳入检索范围')
    
    doc.add_heading('11.3 实用价值', level=3)
    doc.add_paragraph('（1）降低CAD编程门槛：用户无需学习LISP语法即可快速生成代码')
    doc.add_paragraph('（2）提高工作效率：智能推荐相关命令，减少查找时间')
    doc.add_paragraph('（3）知识积累：用户代码保存机制实现个人知识库积累')
    doc.add_paragraph('（4）灵活部署：支持单机和局域网多种部署模式')
    
    doc.add_page_break()
    
    # ========================================
    # 第三部分：代码说明书
    # ========================================
    heading3 = doc.add_heading('', level=1)
    run3 = heading3.add_run('代码说明书')
    run3.font.size = Pt(18)
    run3.font.bold = True
    
    doc.add_heading('一、代码结构', level=2)
    
    # 代码结构表格
    table2 = doc.add_table(rows=15, cols=3)
    table2.style = 'Table Grid'
    
    code_structure = [
        ('序号', '文件名称', '功能说明'),
        ('1', 'main_gui_cloud.py', '客户端主程序，Tkinter GUI界面，包含主窗口、各类输入输出组件'),
        ('2', 'cloud_client.py', '客户端与服务端通信模块，提供HTTP请求封装和缓存管理'),
        ('3', 'kimi_browser.py', 'Kimi浏览器自动化模块，使用Playwright实现浏览器控制'),
        ('4', 'cad_connector.py', 'CAD连接与代码执行模块，通过PythonWin操作CAD'),
        ('5', 'server/cloud_server_rag.py', '服务端主程序，Flask API，包含RAG检索、文件监控等核心逻辑'),
        ('6', 'server/requirements.txt', 'Python依赖清单'),
        ('7', 'start_client.bat', '客户端启动脚本'),
        ('8', 'server/start_server_rag.bat', '服务端启动脚本，包含Ollama路径自动检测'),
        ('9', 'server/stop_server.bat', '服务端停止脚本'),
        ('10', 'server/autocad_basic_commands.txt', 'AutoCAD基本命令库'),
        ('11', 'server/lisp_commands.txt', 'LISP命令描述库'),
        ('12', 'server/user_codes/', '用户代码存储目录'),
        ('13', 'server/test_rag.py', 'RAG功能测试脚本'),
        ('14', 'build_package.py', '项目打包脚本'),
    ]
    
    for i, row_data in enumerate(code_structure):
        row = table2.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True
    
    doc.add_heading('二、核心代码模块说明', level=2)
    
    doc.add_heading('2.1 服务端主模块（cloud_server_rag.py）', level=3)
    doc.add_paragraph('【模块功能】')
    doc.add_paragraph('提供RAG检索、LLM推理、用户代码管理等核心服务')
    doc.add_paragraph('')
    doc.add_paragraph('【主要类/函数】')
    doc.add_paragraph('（1）CommandEmbeddings类')
    doc.add_paragraph('    - load_commands()：加载命令库文件')
    doc.add_paragraph('    - create_embeddings()：生成向量嵌入')
    doc.add_paragraph('    - search()：向量相似度搜索')
    doc.add_paragraph('    - rebuild()：重建向量缓存')
    doc.add_paragraph('')
    doc.add_paragraph('（2）FileChangeHandler类')
    doc.add_paragraph('    - on_modified()：文件修改事件处理')
    doc.add_paragraph('    - callback：触发缓存重建')
    doc.add_paragraph('')
    doc.add_paragraph('（3）API路由')
    doc.add_paragraph('    - /api/health：健康检查')
    doc.add_paragraph('    - /api/query：查询接口')
    doc.add_paragraph('    - /api/user_codes/save：保存用户代码')
    doc.add_paragraph('    - /api/user_codes/list：列出用户代码')
    
    doc.add_heading('2.2 客户端主模块（main_gui_cloud.py）', level=3)
    doc.add_paragraph('【模块功能】')
    doc.add_paragraph('Tkinter图形用户界面，处理用户交互')
    doc.add_paragraph('')
    doc.add_paragraph('【主要类/函数】')
    doc.add_paragraph('（1）CADChatGUI类')
    doc.add_paragraph('    - __init__()：初始化界面组件')
    doc.add_paragraph('    - on_query()：处理查询请求')
    doc.add_paragraph('    - on_save_code()：保存代码到服务器')
    doc.add_paragraph('    - on_connect_cad()：连接CAD')
    doc.add_paragraph('    - update_log()：更新日志显示')
    
    doc.add_heading('2.3 通信模块（cloud_client.py）', level=3)
    doc.add_paragraph('【模块功能】')
    doc.add_paragraph('客户端与服务端的HTTP通信')
    doc.add_paragraph('')
    doc.add_paragraph('【主要类/函数】')
    doc.add_paragraph('（1）CloudClient类')
    doc.add_paragraph('    - query()：发送查询请求')
    doc.add_paragraph('    - save_user_code()：保存用户代码')
    doc.add_paragraph('    - get_user_codes()：获取用户代码列表')
    
    doc.add_heading('2.4 Kimi自动化模块（kimi_browser.py）', level=3)
    doc.add_paragraph('【模块功能】')
    doc.add_paragraph('浏览器自动化，实现Kimi代码分析')
    doc.add_paragraph('')
    doc.add_paragraph('【主要类/函数】')
    doc.add_paragraph('（1）KimiBrowser类')
    doc.add_paragraph('    - login()：登录Kimi网站')
    doc.add_paragraph('    - analyze_code()：发送代码分析请求')
    doc.add_paragraph('    - get_result()：获取分析结果')
    doc.add_paragraph('    - close()：关闭浏览器')
    
    doc.add_heading('2.5 CAD连接模块（cad_connector.py）', level=3)
    doc.add_paragraph('【模块功能】')
    doc.add_paragraph('CAD软件连接和LISP代码执行')
    doc.add_paragraph('')
    doc.add_paragraph('【主要类/函数】')
    doc.add_paragraph('（1）CADConnector类')
    doc.add_paragraph('    - connect()：连接CAD应用')
    doc.add_paragraph('    - is_connected()：检查连接状态')
    doc.add_paragraph('    - run_lisp()：执行LISP代码')
    
    doc.add_heading('三、关键算法说明', level=2)
    
    doc.add_heading('3.1 向量检索算法', level=3)
    doc.add_paragraph('（1）文本向量化：使用Ollama的bge-m3模型')
    doc.add_paragraph('    embedding = ollama.embeddings(model="bge-m3", prompt=text)')
    doc.add_paragraph('')
    doc.add_paragraph('（2）相似度计算：使用sklearn的余弦相似度')
    doc.add_paragraph('    from sklearn.metrics.pairwise import cosine_similarity')
    doc.add_paragraph('    similarities = cosine_similarity(query_embedding, command_embeddings)')
    doc.add_paragraph('')
    doc.add_paragraph('（3）排序取TOP-K')
    doc.add_paragraph('    top_indices = np.argsort(similarities)[0][-top_k:][::-1]')
    
    doc.add_heading('3.2 缓存重建算法', level=3)
    doc.add_paragraph('（1）临时文件机制')
    doc.add_paragraph('    temp_file = "embeddings.tmp.npy"')
    doc.add_paragraph('    final_file = "embeddings.npy"')
    doc.add_paragraph('')
    doc.add_paragraph('（2）原子替换')
    doc.add_paragraph('    import shutil')
    doc.add_paragraph('    shutil.move(temp_file, final_file)')
    
    doc.add_heading('3.3 LLM代码生成', level=3)
    doc.add_paragraph('（1）调用Ollama聊天接口')
    doc.add_paragraph('    response = ollama.chat(model="qwen3:1.7b", messages=[...]')
    doc.add_paragraph('')
    doc.add_paragraph('（2）构建提示词')
    doc.add_paragraph('    prompt = f"根据以下CAD命令生成LISP代码：{command_desc}"')
    
    doc.add_page_break()
    
    # ========================================
    # 第四部分：申请人信息
    # ========================================
    heading4 = doc.add_heading('', level=1)
    run4 = heading4.add_run('申请人（著作权人）信息')
    run4.font.size = Pt(18)
    run4.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('姓名/名称：_________________________')
    doc.add_paragraph()
    doc.add_paragraph('身份证号/组织机构代码：_________________________')
    doc.add_paragraph()
    doc.add_paragraph('通讯地址：_________________________')
    doc.add_paragraph()
    doc.add_paragraph('联系电话：_________________________')
    doc.add_paragraph()
    doc.add_paragraph('电子邮箱：_________________________')
    
    doc.add_paragraph()
    doc.add_paragraph('（申请人签字或盖章）')
    doc.add_paragraph()
    doc.add_paragraph('日期：_____年_____月_____日')
    
    doc.add_page_break()
    
    # ========================================
    # 第五部分：代理信息（如委托代理）
    # ========================================
    heading5 = doc.add_heading('', level=1)
    run5 = heading5.add_run('代理委托书（如有）')
    run5.font.size = Pt(18)
    run5.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('委托人（著作权人）：_________________________')
    doc.add_paragraph()
    doc.add_paragraph('受托人（代理机构）：_________________________')
    doc.add_paragraph()
    doc.add_paragraph('委托事项：□办理软件著作权登记  □领取证书')
    doc.add_paragraph()
    doc.add_paragraph('委托人（签字或盖章）：_________________________')
    doc.add_paragraph()
    doc.add_paragraph('日期：_____年_____月_____日')
    
    # 保存文档
    output_path = 'dist/CADChat_软著申请材料.docx'
    os.makedirs('dist', exist_ok=True)
    doc.save(output_path)
    print(f'软著申请材料已生成: {output_path}')
    
    return output_path

if __name__ == '__main__':
    create_software_copyright_doc()
